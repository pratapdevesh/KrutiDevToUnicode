import os
import shutil
import openpyxl
from docx import Document
from converter import KrutiDevConverter

class FileHandler:
    def __init__(self):
        self.converter = KrutiDevConverter()
    
    def process_file(self, input_path, conversion_type="kd_to_unicode"):
        """
        Process a single file based on extension.
        conversion_type: "kd_to_unicode" or "unicode_to_kd"
        Returns: (success_bool, message_str, changes_count_int)
        """
        ext = os.path.splitext(input_path)[1].lower()
        output_path = self._get_output_path(input_path)
        changes_count = 0
        
        try:
            if ext == '.txt':
                changes_count = self._process_txt(input_path, output_path, conversion_type)
            elif ext == '.xlsx':
                changes_count = self._process_xlsx(input_path, output_path, conversion_type)
            elif ext == '.docx':
                changes_count = self._process_docx(input_path, output_path, conversion_type)
            else:
                return False, f"Unsupported format: {ext}", 0
            
            return True, f"Saved to {output_path}", changes_count
        except Exception as e:
            return False, str(e), 0

    def _get_output_path(self, input_path):
        directory, filename = os.path.split(input_path)
        name, ext = os.path.splitext(filename)
        return os.path.join(directory, f"{name}_converted{ext}")

    def _convert_text(self, text, conversion_type):
        """
        Returns (converted_text, was_changed_bool)
        """
        if not text:
            return text, False
        
        original_text = str(text)
        if conversion_type == "kd_to_unicode":
            converted = self.converter.convert_to_unicode(original_text)
        else:
            converted = self.converter.convert_to_krutidev(original_text)
            
        return converted, (converted != original_text)

    def _process_txt(self, input_path, output_path, conversion_type):
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        converted_content, changed = self._convert_text(content, conversion_type)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(converted_content)
            
        return 1 if changed else 0

    def _process_xlsx(self, input_path, output_path, conversion_type):
        shutil.copy2(input_path, output_path)
        changes = 0
        
        wb = openpyxl.load_workbook(output_path)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        if cell.value.startswith('='):
                            continue
                        try:
                            converted, changed = self._convert_text(cell.value, conversion_type)
                            if changed:
                                cell.value = converted
                                changes += 1
                        except Exception:
                            continue
        wb.save(output_path)
        return changes

    def _process_docx(self, input_path, output_path, conversion_type):
        doc = Document(input_path)
        changes = 0
        
        # Process paragraphs
        for para in doc.paragraphs:
            changes += self._process_runs(para.runs, conversion_type)
            
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        changes += self._process_runs(para.runs, conversion_type)
        
        doc.save(output_path)
        return changes
    
    def _process_runs(self, runs, conversion_type):
        local_changes = 0
        for run in runs:
            if run.text:
                converted, changed = self._convert_text(run.text, conversion_type)
                if changed:
                    run.text = converted
                    local_changes += 1
        return local_changes
